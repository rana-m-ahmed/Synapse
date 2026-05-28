"""
Synapse Backend — DOCX Parser
================================
Extracts text from Microsoft Word (.docx) files using python-docx.
Handles both paragraphs and table content.
"""

import io
import logging

from docx import Document

from app.core.exceptions import DocumentProcessingError
from app.document_processing.parsers.base import BaseParser, ParsedDocument

logger = logging.getLogger("synapse.parsers.docx")


class DocxParser(BaseParser):
    """
    Parses DOCX files by extracting paragraphs and table content.
    Tables are converted to readable row-by-row text.
    """

    async def parse(self, file_content: bytes, file_name: str) -> ParsedDocument:
        """
        Extract text from a DOCX file.

        Extracts both paragraph text and table content. Tables are rendered
        as pipe-separated rows for readability in RAG context.

        Args:
            file_content: Raw DOCX bytes.
            file_name: Original filename.

        Returns:
            ParsedDocument with full text and metadata.

        Raises:
            DocumentProcessingError: If the file is corrupt or unreadable.
        """
        try:
            doc = Document(io.BytesIO(file_content))
        except Exception as e:
            logger.error(f"Failed to read DOCX '{file_name}': {e}")
            raise DocumentProcessingError(
                f"Could not read DOCX file '{file_name}'. The file may be corrupt."
            )

        parts: list[str] = []

        # Extract paragraph text
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        # Extract table content
        for table_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):  # Skip empty rows
                    table_rows.append(" | ".join(cells))

            if table_rows:
                parts.append("\n".join(table_rows))

        full_text = "\n\n".join(parts)

        if not full_text.strip():
            raise DocumentProcessingError(
                f"DOCX file '{file_name}' contains no extractable text."
            )

        logger.info(
            f"Parsed DOCX '{file_name}': {len(doc.paragraphs)} paragraphs, "
            f"{len(doc.tables)} tables, {len(full_text)} chars"
        )

        return ParsedDocument(
            text=full_text,
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "source_file": file_name,
            },
        )
